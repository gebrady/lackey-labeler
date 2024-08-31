##### Suite Color Chooser Beta 1.0 #####
# Developed August-September 2022, Designed June 2022-Present
# Written by Graham E. Brady, with conceptual help from Rohan Shaiva and syntax help from Stack Overflow / Python docs
##### Defines Label Maker Class and Main functions when run from terminals #####

# Script to populate a label template using imported data.
# Relevant data from a .csv file will be filled into the standardized pdf format for printing
# This program is specifically for XRF bead bag labels. 
# 

##### Libraries #####

import utm
import pandas as pd
from docx import Document
import json
import os

from docx.shared import RGBColor

class LabelMaker():
    def __init__(self, template, dataframe, colorcodes):
        self.doc = Document('LabelTemplate.docx')
        ##### Prepare template #####
        # doc.tables lists all tables in document
        #template_format = doc.tables[0].cell(0, 0).text
        #print("Data format of template: " + template_format)

        self.template_table = self.doc.tables[0] # if only one table in the template file
        self.template_table.allow_autofit = False

        self.num_rows = len(self.template_table.rows)
        self.num_cols = len(self.template_table.columns)
        self.num_samples = self.num_rows*self.num_cols

        ##### Your data #####

        self.data = dataframe # read data from dataTable
        self.template = template
    ##### Converting GPS data to Decimal Degrees #####

        def convert_GPS(row):
            print(row)
            lat, long = utm.to_latlon(row['Easting'], row['Northing'], 11, northern=True)
            return pd.Series({"Lat": lat, "Long": long})

        #slate = slate.merge(jrd.apply(convert_GPS, axis=1), left_index= True, right_index= True) 

        # Samples are now subselected and contain all template info

        ##### Converting data to the template format and file #####
        suite = 'Suite'
        sample = 'Sample'
        lith = 'Lithology'
        lat = 'Latitude'
        long = 'Longitude'
        misc = 'Misc'
        PI = 'Collector'
        self.colnames = [suite, sample, lith, lat, long, misc, PI]  #Consolidate these column names for use in the formatting helper function

        self.df = self.data[['Suite', 'Sample', 'Lithology', 'Latitude', 'Longitude', 'Misc', 'Collector']]
        self.df.columns = ['suite', 'sample', 'lith', 'lat', 'long', 'misc','PI']

        for i in ['lat', 'long']:
            self.df[i] = self.df[i].round(5).astype(str)

        self.df['latlong'] = self.df['lat'] + '°, '+self.df['long'] + '°' #set up single datum for latlong

        self.df['misc2'] = self.df['misc'] + '; '+self.df['PI'] #set up single datum for latlong

        self.df = self.df[['suite','sample','lith','latlong','misc2']]

        self.count_rows_cell = 5 # count of rows to copy into each cell of the template. (7-2 for missing individual lat and long)

        self.colorDF = pd.DataFrame(colorcodes, columns = ['suite','rgb'])

        def format_fill_data(data):
            """converts a sample's data to the template format for input."""
            out = str(data[0])
            for e in data[1:]:
                out += '\n'  + str(e) 
            # 'suite\nsample\nlith\nage\nlat\nlong\nmisc'
            return out

    # Converting data cols to the relevant headers for each sample,
                    # which on the template appear as rows

    #data_rows = ['Suite', 'Sample_ID', 'Lithology', 'Age', 
    #             'Latitude', 'Longitude', 'Year/PI/Study']


    #df.columns = ['suite', 'sample', 'lith', 'latlong', 'misc', 'lat', 'long']
    """
    self.suites_colors = df.suite.unique().tolist()
    print('Suites are: ', suites_colors)
    """
    #colorcodes = [('Suite 1', (0, 142, 0)), ('Suite 2', (148, 22, 81)), ('Suite 3', (255, 38, 0))]
    #colorcodes = eval(input('Copy and paste your color code from Tkinter.'))

    def writeTable(self, filename):
        for e in range(self.num_samples):
            if e >= len(self.df):
                break
            print(e)
            self.template_table.cell(0,e).text = '' #clear existing data from template
            data = self.df.iloc[e]
            print(data)
            styles_xrf = ['Suite', 'Sample', 'Lithology', 'LatLong', 'Misc']
            print(styles_xrf)
            for i in range(self.count_rows_cell):
                print(i)
                p = self.template_table.cell(0,e).paragraphs[i]
                r = p.add_run(data[i])
                print('adding data: ', data[i])
                if styles_xrf[i] == 'Suite':
                    print('changing color')
                    r.font.color.rgb = RGBColor(self.colorDF['rgb'][self.colorDF['suite']==data.suite].values[0][0],
                                                    self.colorDF['rgb'][self.colorDF['suite']==data.suite].values[0][1],
                                                    self.colorDF['rgb'][self.colorDF['suite']==data.suite].values[0][2])
                    print('color changed')
                else: r.font.color.rgb = RGBColor(0,0,0)
                p.style = self.doc.styles[styles_xrf[i]]
                if i == self.count_rows_cell - 1: pass
                else: p = self.template_table.cell(0,e).add_paragraph()
                print('row written:', str(data[i]))
        print('writing document')
        self.doc.save(filename + "_Labels.docx")
        print('document written')

if __name__ == "__main__":
    lm = LabelMaker()
    lm.writeTable('test')
